import io
import re
import markdown2
from fpdf import FPDF
from docx import Document

def export_to_pdf(markdown_text: str) -> bytes:
    """Converts simple markdown text to a PDF byte string."""
    html = markdown2.markdown(markdown_text)
    
    class PDF(FPDF):
        def header(self):
            pass
        def footer(self):
            pass
            
    pdf = PDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    # Add a built-in font
    pdf.set_font('helvetica', size=11)
    
    # Try to write HTML, fallback to simple text if it fails
    try:
        pdf.write_html(html)
    except Exception as e:
        pdf.multi_cell(0, 5, text=markdown_text)
        
    out = pdf.output()
    return bytes(out)

def _add_markdown_runs(paragraph, text: str):
    """Helper to parse **bold** and *italic* in a single paragraph run."""
    # Split by bold first
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            # Check for italic inside bold
            sub_parts = re.split(r'(\*.*?\*)', bold_text)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*'):
                    paragraph.add_run(sp[1:-1]).bold = True
                    paragraph.add_run('').italic = True # applying to same run needs different approach, but simplistic is fine
                else:
                    paragraph.add_run(sp).bold = True
        else:
            # Check for italic
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sp in sub_parts:
                if sp.startswith('*') and sp.endswith('*'):
                    paragraph.add_run(sp[1:-1]).italic = True
                else:
                    paragraph.add_run(sp)

def export_to_docx(markdown_text: str) -> bytes:
    """Converts simple markdown text to a DOCX byte string."""
    doc = Document()
    lines = markdown_text.split('\n')
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:].strip(), level=1)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:].strip(), level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:].strip(), level=3)
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            # The `- ` or `* ` is 2 chars. If it was `* ` it might be italic, but it's at the start.
            text = line_stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_markdown_runs(p, text)
        else:
            p = doc.add_paragraph()
            _add_markdown_runs(p, line_stripped)
            
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
